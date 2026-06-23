import io
import json
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from app.core.dependencies import get_current_user
from app.db.session import get_db
from app.models.audit_report import AuditReport
from app.models.user import User

reportlab_installed = True
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.colors import HexColor
    from reportlab.lib import colors
except ImportError as e:
    import logging
    logging.getLogger(__name__).error(f"Failed to import reportlab: {e}")
    reportlab_installed = False

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    pass

router = APIRouter(prefix="/reports", tags=["Export"])

@router.get("/{report_id}/export/pdf", summary="Export Compliance Report as PDF")
def export_pdf(
    report_id: int,
    db: Session = Depends(get_db),
    _user: User = Depends(get_current_user),
):
    """
    Generate a formatted PDF for the given AuditReport ID.
    Returns the file as an application/pdf attachment.
    """
    if not reportlab_installed:
        raise HTTPException(status_code=500, detail="PDF generation library (reportlab) not installed. Please install required system libraries (e.g. libfreetype6).")
        
    report = db.query(AuditReport).filter(AuditReport.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
        styles = getSampleStyleSheet()
        
        # Add custom styles
        styles['Title'].textColor = HexColor('#1e1b4b')
        styles['Title'].fontSize = 22
        styles['Title'].leading = 26
        styles['Title'].spaceAfter = 15
        
        styles.add(ParagraphStyle(
            name='ReportSubTitle', 
            parent=styles['Heading2'], 
            textColor=HexColor('#4f46e5'),
            fontSize=14,
            leading=18,
            spaceBefore=14,
            spaceAfter=8,
            keepWithNext=True
        ))
        
        styles.add(ParagraphStyle(
            name='Body',
            parent=styles['Normal'],
            fontSize=10,
            leading=14,
            textColor=HexColor('#374151'),
            spaceAfter=8
        ))
        
        styles.add(ParagraphStyle(
            name='FindingBullet',
            parent=styles['Normal'],
            fontSize=9.5,
            leading=13.5,
            textColor=HexColor('#4b5563'),
            leftIndent=15,
            firstLineIndent=-10,
            spaceAfter=6
        ))
        
        styles.add(ParagraphStyle(name='RiskHigh', parent=styles['Normal'], textColor=HexColor('#dc2626'), fontName='Helvetica-Bold'))
        styles.add(ParagraphStyle(name='RiskMedium', parent=styles['Normal'], textColor=HexColor('#d97706'), fontName='Helvetica-Bold'))
        styles.add(ParagraphStyle(name='RiskLow', parent=styles['Normal'], textColor=HexColor('#16a34a'), fontName='Helvetica-Bold'))

        story = []
        
        # Title
        story.append(Paragraph(f"Compliance Audit Report #{report.id}", styles['Title']))
        
        # Divider Line
        divider = Table([[""]], colWidths=[504])
        divider.setStyle(TableStyle([
            ('LINEBELOW', (0,0), (-1,-1), 1.5, HexColor('#6366f1')),
            ('BOTTOMPADDING', (0,0), (-1,-1), 0),
            ('TOPPADDING', (0,0), (-1,-1), 0),
        ]))
        story.append(divider)
        story.append(Spacer(1, 15))
        
        # Executive Summary
        story.append(Paragraph("Executive Summary", styles['ReportSubTitle']))
        summary_text = (
            f"This compliance report was compiled by the Enterprise AI Platform. "
            f"An evaluation of company policies against regulatory frameworks resulted in a compliance score of {report.compliance_score}%, "
            f"indicating a {report.risk.upper()} risk level. A total of {report.violation_count} violations were detected. "
            f"Please review the detailed findings and implement the recommendations listed below to achieve full compliance."
        )
        story.append(Paragraph(summary_text, styles['Body']))
        story.append(Spacer(1, 8))

        # Metrics Table
        story.append(Paragraph("Audit Metrics", styles['ReportSubTitle']))
        
        risk_style = styles[f'Risk{report.risk}'] if f'Risk{report.risk}' in styles else styles['Normal']
        
        metrics_data = [
            [Paragraph("<b>Audit Metric</b>", styles['Body']), Paragraph("<b>Value / Status</b>", styles['Body'])],
            [Paragraph("Compliance Score", styles['Body']), Paragraph(f"<b>{report.compliance_score}%</b>", styles['Body'])],
            [Paragraph("Risk Level", styles['Body']), Paragraph(f"<b>{report.risk.upper()}</b>", risk_style)],
            [Paragraph("Violations Detected", styles['Body']), Paragraph(f"<b>{report.violation_count}</b>", styles['Body'])],
            [Paragraph("Audit Timestamp", styles['Body']), Paragraph(report.audit_timestamp, styles['Body'])],
            [Paragraph("Auditor", styles['Body']), Paragraph(report.auditor, styles['Body'])],
        ]
        metrics_table = Table(metrics_data, colWidths=[200, 304])
        metrics_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), HexColor('#f3f4f6')),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('GRID', (0,0), (-1,-1), 0.5, HexColor('#e5e7eb')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, HexColor('#f9fafb')]),
            ('TOPPADDING', (0,0), (-1,-1), 6),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ]))
        story.append(metrics_table)
        story.append(Spacer(1, 12))

        # Compliance Findings
        story.append(Paragraph("Compliance Findings", styles['ReportSubTitle']))
        findings = json.loads(report.issues) if report.issues else []
        if findings:
            for f in findings:
                story.append(Paragraph(f"&bull; {f}", styles['FindingBullet']))
        else:
            story.append(Paragraph("No compliance violations detected.", styles['Body']))
        story.append(Spacer(1, 8))

        # Recommendations
        story.append(Paragraph("Remediation Recommendations", styles['ReportSubTitle']))
        recs = json.loads(report.recommendations) if report.recommendations else []
        if recs:
            for r in recs:
                story.append(Paragraph(f"&bull; {r}", styles['FindingBullet']))
        else:
            story.append(Paragraph("No actionable recommendations required.", styles['Body']))

        doc.build(story)
        buffer.seek(0)
        
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=Compliance_Report_{report.id}.pdf"}
        )
    except NameError as exc:
        raise HTTPException(status_code=500, detail=f"NameError: {exc}")
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {exc}")


@router.get("/{report_id}/export/docx", summary="Export Compliance Report as DOCX")
def export_docx(
    report_id: int,
    db: Session = Depends(get_db),
    _user: User = Depends(get_current_user),
):
    report = db.query(AuditReport).filter(AuditReport.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    try:
        doc = Document()
        
        # Title
        title = doc.add_heading(f"Compliance Audit Report #{report.id}", 0)
        title.style.font.color.rgb = None  # Use word default/standard navy
        
        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        summary_text = (
            f"This compliance report was compiled by the Enterprise AI Platform. "
            f"An evaluation of company policies against regulatory frameworks resulted in a compliance score of {report.compliance_score}%, "
            f"indicating a {report.risk.upper()} risk level. A total of {report.violation_count} violations were detected. "
            f"Please review the detailed findings and implement the recommendations listed below to achieve full compliance."
        )
        doc.add_paragraph(summary_text)

        # Metrics Table
        doc.add_heading("Audit Metrics", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Audit Metric'
        hdr_cells[1].text = 'Value / Status'
        
        metrics = [
            ("Compliance Score", f"{report.compliance_score}%"),
            ("Risk Level", report.risk.upper()),
            ("Violations Detected", str(report.violation_count)),
            ("Audit Timestamp", report.audit_timestamp),
            ("Auditor", report.auditor)
        ]
        for metric, val in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = val
            
        doc.add_paragraph()  # Spacing

        # Compliance Findings
        doc.add_heading("Compliance Findings", level=1)
        findings = json.loads(report.issues) if report.issues else []
        if findings:
            for f in findings:
                doc.add_paragraph(f, style='List Bullet')
        else:
            doc.add_paragraph("No compliance violations detected.")

        # Recommendations
        doc.add_heading("Remediation Recommendations", level=1)
        recs = json.loads(report.recommendations) if report.recommendations else []
        if recs:
            for r in recs:
                doc.add_paragraph(r, style='List Bullet')
        else:
            doc.add_paragraph("No actionable recommendations required.")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=Compliance_Report_{report.id}.docx"}
        )
    except NameError:
        raise HTTPException(status_code=500, detail="DOCX generation library (python-docx) not installed")
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {exc}")
